#!/usr/bin/env python3
"""
Simple DOCX to PDF Converter
---------------------------
This script converts DOCX files to PDF while preserving formatting, tables, and images.
Just set your input and output directories below, and run the script.

Requirements:
    - docx2pdf (handles the conversion with high fidelity)
    - python-docx (for reading DOCX metadata)
    - tqdm (for progress bar)

Install dependencies:
    pip install docx2pdf python-docx tqdm
"""

import os
import sys
import time
from pathlib import Path
from tqdm import tqdm
import logging
import platform

# ====== CONFIGURATION - CHANGE THESE VALUES ======
# Directory containing DOCX files to convert
INPUT_DIRECTORY = r"/var/home/noor/D/University_projects/computerNetwork/cn_u2210167.docx"  # Change this to your input directory

# Directory where PDF files will be saved
OUTPUT_DIRECTORY = r"/var/home/noor/D/University_projects/computerNetwork/CN_MajorAssignment_U2210167.pdf"  # Change this to your output directory

# Process files in subdirectories
PROCESS_SUBDIRECTORIES = True

# Keep detailed log
ENABLE_LOGGING = True
# =================================================

# Configure logging
if ENABLE_LOGGING:
    logging.basicConfig(
        level=logging.INFO,
        format="%(asctime)s [%(levelname)s] %(message)s",
        handlers=[
            logging.FileHandler("docx_to_pdf_conversion.log"),
            logging.StreamHandler(sys.stdout)
        ]
    )
else:
    logging.basicConfig(
        level=logging.INFO,
        format="%(asctime)s [%(levelname)s] %(message)s",
        handlers=[
            logging.StreamHandler(sys.stdout)
        ]
    )
logger = logging.getLogger(__name__)

def check_requirements():
    """Check if required dependencies are installed"""
    try:
        import docx2pdf
        import docx
        import tqdm
    except ImportError as e:
        module = str(e).split("'")[1]
        logger.error(f"Required module not found: {module}")
        logger.error("Please install required dependencies with: pip install docx2pdf python-docx tqdm")
        return False
    
    # On non-Windows platforms, check if LibreOffice is installed
    if platform.system() != "Windows":
        # Try to find LibreOffice executable
        libreoffice_found = False
        for cmd in ["libreoffice", "soffice"]:
            result = os.system(f"which {cmd} > /dev/null 2>&1")
            if result == 0:
                libreoffice_found = True
                break
        
        if not libreoffice_found:
            logger.error("LibreOffice not found. It's required for PDF conversion on non-Windows platforms.")
            logger.error("Please install LibreOffice and try again.")
            return False
    
    return True

def get_doc_info(docx_path):
    """Extract metadata from DOCX file"""
    try:
        import docx
        doc = docx.Document(docx_path)
        core_properties = doc.core_properties
        title = core_properties.title or "Untitled"
        author = core_properties.author or "Unknown"
        created = core_properties.created or "Unknown"
        
        # Count tables and images
        table_count = len(doc.tables)
        image_count = 0
        
        # Count images (needs to iterate through all runs in all paragraphs)
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                image_count += len(run._element.findall('.//a:blip', 
                                  namespaces=docx.oxml.ns.nsmap))
                
        return {
            "title": title,
            "author": author,
            "created": created,
            "tables": table_count,
            "images": image_count
        }
    except Exception as e:
        logger.warning(f"Couldn't extract document info: {str(e)}")
        return {
            "title": "Unknown",
            "author": "Unknown",
            "created": "Unknown",
            "tables": "Unknown", 
            "images": "Unknown"
        }

def convert_file(input_path, output_path=None):
    """
    Convert a single DOCX file to PDF
    
    Args:
        input_path (str): Path to the DOCX file
        output_path (str, optional): Path for the output PDF file
    """
    from docx2pdf import convert
    
    input_path = Path(input_path).resolve()
    
    if not input_path.exists():
        logger.error(f"File not found: {input_path}")
        return False
    
    if input_path.suffix.lower() != '.docx':
        logger.error(f"Not a DOCX file: {input_path}")
        return False
    
    # If output path is not specified, use same name but with PDF extension
    if not output_path:
        output_path = input_path.with_suffix('.pdf')
    else:
        output_path = Path(output_path).resolve()
    
    # Extract document info before conversion
    doc_info = get_doc_info(input_path)
    logger.info(f"Converting: {input_path.name}")
    logger.info(f"Document info: Title: {doc_info['title']}, Author: {doc_info['author']}")
    logger.info(f"Contains {doc_info['tables']} tables and {doc_info['images']} images")
    
    try:
        # Start timing
        start_time = time.time()
        
        # Convert the file
        # On Windows, docx2pdf uses MS Word in the background
        # On other platforms, it uses LibreOffice
        convert(input_path, output_path)
        
        duration = time.time() - start_time
        logger.info(f"✅ Conversion successful: {input_path.name} → {output_path.name} ({duration:.2f}s)")
        
        # Check output file size
        if output_path.exists():
            input_size = input_path.stat().st_size / (1024 * 1024)  # MB
            output_size = output_path.stat().st_size / (1024 * 1024)  # MB
            logger.info(f"File sizes: Input: {input_size:.2f}MB, Output: {output_size:.2f}MB")
            return True
        else:
            logger.error(f"❌ Output file not created: {output_path}")
            return False
            
    except Exception as e:
        logger.error(f"❌ Conversion failed for {input_path.name}: {str(e)}")
        return False

def convert_directory(input_dir, output_dir, recursive=False):
    """
    Convert all DOCX files in a directory to PDF
    
    Args:
        input_dir (str): Directory containing DOCX files
        output_dir (str): Directory for output PDF files
        recursive (bool): Whether to search for DOCX files in subdirectories
    """
    input_dir = Path(input_dir).resolve()
    
    if not input_dir.is_dir():
        logger.error(f"Not a directory: {input_dir}")
        return False
    
    output_dir = Path(output_dir).resolve()
    # Create output directory if it doesn't exist
    output_dir.mkdir(parents=True, exist_ok=True)
    
    # Find all DOCX files
    if recursive:
        docx_files = list(input_dir.glob('**/*.docx'))
    else:
        docx_files = list(input_dir.glob('*.docx'))
    
    if not docx_files:
        logger.warning(f"No DOCX files found in {input_dir}")
        return False
    
    logger.info(f"Found {len(docx_files)} DOCX files in {input_dir}")
    
    # Convert each file
    success_count = 0
    for docx_file in tqdm(docx_files, desc="Converting files"):
        # Maintain relative path structure for recursive conversion
        if recursive:
            rel_path = docx_file.relative_to(input_dir)
            output_path = output_dir / rel_path.with_suffix('.pdf')
            # Create subdirectory in output_dir if needed
            output_path.parent.mkdir(parents=True, exist_ok=True)
        else:
            output_path = output_dir / docx_file.with_suffix('.pdf').name
        
        if convert_file(docx_file, output_path):
            success_count += 1
    
    logger.info(f"✅ Conversion complete: {success_count}/{len(docx_files)} files converted successfully")
    return success_count > 0

def main():
    print("=" * 60)
    print("DOCX to PDF Converter - Starting")
    print("=" * 60)
    
    # Check if required modules are installed
    if not check_requirements():
        print("Missing required modules. Please install with:")
        print("pip install docx2pdf python-docx tqdm")
        return
    
    # Print system info
    print(f"System: {platform.system()} {platform.version()}")
    
    # Check if input directory exists
    input_dir = Path(INPUT_DIRECTORY)
    if not input_dir.exists() or not input_dir.is_dir():
        print(f"Error: Input directory does not exist: {INPUT_DIRECTORY}")
        return
    
    # Create output directory if needed
    output_dir = Path(OUTPUT_DIRECTORY)
    if not output_dir.exists():
        print(f"Creating output directory: {OUTPUT_DIRECTORY}")
        output_dir.mkdir(parents=True, exist_ok=True)
    
    # Convert files
    print(f"Input directory: {INPUT_DIRECTORY}")
    print(f"Output directory: {OUTPUT_DIRECTORY}")
    print(f"Process subdirectories: {'Yes' if PROCESS_SUBDIRECTORIES else 'No'}")
    print("-" * 60)
    
    convert_directory(input_dir, output_dir, PROCESS_SUBDIRECTORIES)
    
    print("=" * 60)
    print("Conversion process complete!")
    print("=" * 60)

if __name__ == "__main__":
    main()
